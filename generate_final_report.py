import sys
import os
from pathlib import Path

# Add the temp site-packages to sys.path
temp_site_packages = "/home/pratyush/.gemini/tmp/9c760671870f7b4f11154a118c7345dfba3ae8c8edb9007ced8ed74c47189df0/site-packages"
if temp_site_packages not in sys.path:
    sys.path.append(temp_site_packages)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError as e:
    print(f"Error importing python-docx: {e}")
    print(f"sys.path: {sys.path}")
    sys.exit(1)

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Emotion Classifier Using Verilog', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Final Project Report', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Date: 27 November 2025').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    p = doc.add_paragraph(
        "This project implements a hardware/software co-design for real-time emotion classification. "
        "It combines a high-performance face detector implemented in Verilog (for FPGA deployment) "
        "with a flexible emotion classification neural network running in Python. The system leverages "
        "the Verilog Procedural Interface (VPI) to enable seamless communication between the hardware "
        "simulation and the software model."
    )

    # 2. System Architecture
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph(
        "The system consists of three main layers: the Verilog Hardware Layer, the VPI Interface Layer, "
        "and the Python AI Server Layer."
    )
    
    # Architecture Diagram (ASCII style in Monospace)
    doc.add_heading('2.1 Architecture Diagram', level=2)
    diagram = """
┌─────────────────────────────────────────────────────────────────────────────┐
│                           VERILOG HARDWARE LAYER                            │
│                         (Icarus Verilog Simulator)                          │
├─────────────────────────────────────────────────────────────────────────────┤
│  ┌───────────────────────────────────────────────────────────────────────┐  │
│  │                        Face Detector Module                           │  │
│  │                      (Haar Cascade Algorithm)                         │  │
│  └───────────────────────────────────────────────────────────────────────┘  │
│                                     │                                       │
└─────────────────────────────────────┼─────────────────────────────────────┘
                                      ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│                         VPI C INTERFACE LAYER                               │
│                    (verilog_python_interface.c)                             │
└─────────────────────────────────────┼─────────────────────────────────────┘
                                      ▼
┌─────────────────────────────────────────────────────────────────────────────┐
│                         PYTHON AI SERVER LAYER                              │
│                          (emotion_server.py)                                │
│  ┌───────────────────────────────────────────────────────────────────────┐  │
│  │                    Mini-Xception Neural Network                       │  │
│  └───────────────────────────────────────────────────────────────────────┘  │
└─────────────────────────────────────────────────────────────────────────────┘
    """
    paragraph = doc.add_paragraph(diagram)
    paragraph.style.font.name = 'Courier New'
    paragraph.style.font.size = Pt(8)

    # Component descriptions
    doc.add_heading('2.2 Key Components', level=2)
    
    components = [
        ("Face Detector (Verilog)", "Implements the Viola-Jones Haar Cascade algorithm. It processes 64x64 grayscale images using integral images and a cascade of weak classifiers to detect faces."),
        ("VPI Interface (C/C++)", "Acts as a bridge, extracting the Region of Interest (ROI) from the Verilog memory and transmitting it over a TCP socket."),
        ("Emotion Server (Python)", "Hosts a Mini-Xception Convolutional Neural Network (CNN). It receives the image data, preprocesses it (resize to 48x48, normalize), and classifies it into one of 7 emotions (Angry, Disgust, Fear, Happy, Sad, Surprise, Neutral).")
    ]
    
    for name, desc in components:
        p = doc.add_paragraph()
        run = p.add_run(name + ": ")
        run.bold = True
        p.add_run(desc)

    # 3. Work Accomplished
    doc.add_heading('3. Work Accomplished', level=1)
    
    work_items = [
        "Implemented the full Haar Cascade detection pipeline in Verilog (Integral Image, Feature Calculation, Strong/Weak Classifiers).",
        "Developed a Python TCP server to host the emotion classification model.",
        "Created a C-based VPI module to facilitate data transfer between Verilog simulation and Python.",
        "Designed a testing pipeline (`test_pipeline.py`) that automates image preparation, simulation, and result verification.",
        "Provided Tcl scripts for seamless integration with Xilinx Vivado for FPGA synthesis."
    ]
    
    for item in work_items:
        doc.add_paragraph(item, style='List Bullet')

    # 4. Findings & Results
    doc.add_heading('4. Findings & Results', level=1)
    doc.add_paragraph(
        "The co-simulation successfully demonstrates the feasibility of offloading the computationally intensive "
        "face detection task to hardware while keeping the complex deep learning inference in software. "
        "This approach balances performance with flexibility."
    )
    
    doc.add_heading('4.1 Latency Analysis (Estimated)', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Stage'
    hdr_cells[1].text = 'Time (ms)'
    hdr_cells[2].text = 'Percentage'
    
    data = [
        ('Face Detection (HW)', '5.0', '50%'),
        ('Data Transfer (VPI/TCP)', '1.0', '10%'),
        ('Emotion Inference (SW)', '3.0', '30%'),
        ('Overhead', '1.0', '10%'),
    ]
    
    for stage, time, pct in data:
        row_cells = table.add_row().cells
        row_cells[0].text = stage
        row_cells[1].text = time
        row_cells[2].text = pct

    # 5. Instructions for Generating Graphs
    doc.add_heading('5. Visualizations & Graphs', level=1)
    doc.add_paragraph(
        "To generate the visual artifacts for this project, follow these instructions:"
    )
    
    doc.add_heading('5.1 Architecture Diagram', level=2)
    doc.add_paragraph("Run the visualization script to see the detailed ASCII block diagram:")
    doc.add_paragraph("python3 verilog_face_detector/visualize_architecture.py", style='Quote')

    doc.add_heading('5.2 Simulation Waveforms', level=2)
    doc.add_paragraph(
        "To view the signal-level behavior of the hardware face detector:"
    )
    steps = [
        "Navigate to the 'verilog_face_detector' directory.",
        "Run 'make run-cosim' to generate the waveform dump.",
        "Run 'make wave' to open GTKWave with the pre-configured 'sim/waveform.gtkw' file."
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')
        
    # Save
    output_path = "Final_Report.docx"
    doc.save(output_path)
    print(f"Report saved to {output_path}")

if __name__ == "__main__":
    create_report()
